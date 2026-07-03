from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rowen_Infante_CV.docx"

BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(70, 70, 70)


def set_run(run, size=11, bold=False, italic=False, color=BLACK, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", size=11, bold=False, italic=False, after=6, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text):
    p = paragraph(doc, text.upper(), size=12, bold=True, after=4, before=12)
    p.style = doc.styles["Heading 1"]
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=10.8)
    return p


def role(doc, title, company, dates=""):
    p = paragraph(doc, after=2, before=6)
    set_run(p.add_run(title), size=11, bold=True)
    set_run(p.add_run(f" | {company}"), size=11)
    if dates:
        set_run(p.add_run(f" | {dates}"), size=11, color=MUTED)
    return p


def create_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = BLACK
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)

    name = paragraph(doc, "ROWEN INFANTE", size=20, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    name.paragraph_format.space_before = Pt(0)
    paragraph(
        doc,
        "SAP BTP Administrator | SAP HANA | SAP Generative AI | Web Development",
        size=11,
        after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    paragraph(
        doc,
        "rowen.infante@gmail.com | github.com/rowenautomation | linkedin.com/in/rowen-infante-539279210",
        size=10,
        after=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    heading(doc, "Professional Summary")
    paragraph(
        doc,
        "SAP industry professional at Accenture with experience in SAP BTP administration, SAP BTP integration, SAP HANA, "
        "SAP Generative AI, identity and access management, roles, role collections, transport management, CI/CD, and "
        "SAP Build Work Zone. Background includes Manufacturing Engineering from Mapua University, process engineering "
        "experience, web development, GitHub-based project work, AI-assisted web development, and AI-assisted video creation.",
        after=6,
    )

    heading(doc, "Core Skills")
    paragraph(
        doc,
        "SAP BTP, SAP BTP Integration, roles, role collections, identity, access management, SAP CTMS, SAP CI/CD service, "
        "SAP Build Work Zone, SAP HANA, SAP Generative AI, GitHub, Next.js, Supabase, ecommerce UI, portfolio development, "
        "HANA Studio, AI-assisted web development, AI-assisted video creation, short-form content experiments, visual concepts, "
        "Higgsfield, OpenArt, Claude, ChatGPT, Gemini, Flow, data analytics foundations.",
        after=6,
    )

    heading(doc, "Professional Experience")
    role(doc, "SAP BTP Administrator / SAP Industry Professional", "Accenture", "Current")
    bullet(doc, "Develop and support SAP BTP administration workflows for enterprise SAP environments.")
    bullet(doc, "Support SAP BTP integration, platform administration, roles, role collections, identity, and access management requirements.")
    bullet(doc, "Work with SAP CTMS, SAP CI/CD service, SAP Build Work Zone, SAP HANA, and SAP Generative AI.")
    bullet(doc, "Use GitHub and modern web development practices while building portfolio and ecommerce projects.")

    role(doc, "Process Engineer", "SFA Semicon Philippines Corporation", "June 2022 - December 2022")
    bullet(doc, "Applied engineering discipline, process analysis, and production support in a manufacturing environment.")
    bullet(doc, "Strengthened practical problem-solving, documentation, teamwork, and process improvement skills.")

    heading(doc, "Projects")
    role(doc, "Obsidian - Skincare Ecommerce Website", "Personal Project", "Under Development")
    bullet(doc, "Built a luxury skincare ecommerce website concept with product sections, brand storytelling, reviews, cart/search navigation, and an about page.")
    bullet(doc, "Used Next.js, Supabase, ecommerce interface design, HANA Studio awareness, AI-assisted web development, and GitHub-based development workflow.")

    role(doc, "Portfolio Website", "Personal Project", "2026")
    bullet(doc, "Created a personal portfolio for SAP BTP work, certifications, web projects, and AI creative experiments.")
    bullet(doc, "Implemented a black-and-white portfolio UI inspired by clean developer portfolio patterns.")

    role(doc, "AI Video Creation", "Creative AI Project", "Ongoing")
    bullet(doc, "Create AI-assisted short-form videos using Higgsfield, OpenArt, Claude, ChatGPT, Gemini, Flow, visual concepts, motion, and editing workflows.")

    heading(doc, "Certifications")
    bullet(doc, "SAP Certified - SAP BTP Administrator")
    bullet(doc, "SAP Certified - SAP Generative AI Developer")
    bullet(doc, "Foundations: Data, Data, Everywhere - Coursera / Google")

    heading(doc, "Education")
    role(doc, "Bachelor of Science in Manufacturing Engineering", "Mapua University", "June 2015 - August 2021")

    heading(doc, "Languages")
    paragraph(doc, "English | Filipino", after=0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    create_doc()
