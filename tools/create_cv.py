from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Rowen_Infante_CV.docx"

BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(28, 28, 28)
MUTED = RGBColor(95, 95, 95)
FAINT = RGBColor(150, 150, 150)
LINE = "D9D9D9"
FILL = "F7F7F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_run(run, size=10.5, color=BLACK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def para(doc, text="", size=10.5, color=BLACK, bold=False, italic=False, after=4, before=0, align=None, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 11.5, 3: 10.5}
    afters = {1: 5, 2: 4, 3: 2}
    p = para(doc, text.upper(), size=sizes[level], color=DARK, bold=True, after=afters[level], before=10 if level == 1 else 5)
    p.style = doc.styles[f"Heading {level}"]
    return p


def bullet(doc, text, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run(run, size=10.2, color=BLACK)
    return p


def add_rule(paragraph, color=LINE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def section_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    for row_idx, (label, value) in enumerate(rows):
        row = table.rows[row_idx]
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(4.95)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=70, bottom=70, start=90, end=90)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LINE})
        p_label = row.cells[0].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(0)
        set_run(p_label.add_run(label), size=9.2, color=FAINT, bold=True)
        p_value = row.cells[1].paragraphs[0]
        p_value.paragraph_format.space_after = Pt(0)
        set_run(p_value.add_run(value), size=10.3, color=BLACK)
    return table


def create_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.1

    for idx, size in [(1, 14), (2, 11.5), (3, 10.5)]:
        st = styles[f"Heading {idx}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = DARK
        st.paragraph_format.space_before = Pt(9 if idx == 1 else 5)
        st.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("Rowen Infante | CV"), size=8.5, color=FAINT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("rowen.infante@gmail.com | github.com/rowenautomation | linkedin.com/in/rowen-infante-539279210"), size=8.5, color=FAINT)

    title = para(doc, "ROWEN INFANTE", size=23, color=BLACK, bold=True, after=2)
    subtitle = para(doc, "SAP BTP Administrator | SAP HANA | SAP Generative AI | Web Development", size=11.5, color=MUTED, after=8)
    contact = para(doc, "rowen.infante@gmail.com  |  github.com/rowenautomation  |  linkedin.com/in/rowen-infante-539279210", size=9.7, color=MUTED, after=8)
    rule = para(doc, "", after=8)
    add_rule(rule)

    heading(doc, "Professional Profile", 1)
    para(
        doc,
        "SAP industry professional at Accenture with a Manufacturing Engineering foundation and process engineering experience. "
        "Focused on SAP BTP administration, identity and access setup, destinations, CTMS, CI/CD, SAP Build Work Zone, SAP HANA, and SAP Generative AI. "
        "Also builds web projects and explores AI-assisted video creation as a creative technical practice.",
        size=10.5,
        color=BLACK,
        after=6,
    )

    heading(doc, "Core Skills", 1)
    section_table(
        doc,
        [
            ("SAP BTP", "Destination configuration and creation, roles, role collections, identity, access management"),
            ("SAP Services", "SAP CTMS, SAP CI/CD service, SAP Build Work Zone, SAP HANA, SAP Generative AI"),
            ("Web & Tools", "GitHub, Next.js, Supabase, ecommerce UI, portfolio development"),
            ("Creative AI", "AI-assisted video creation, short-form content experiments, visual concepts"),
        ],
    )

    heading(doc, "Experience", 1)
    heading(doc, "SAP BTP Administrator / SAP Industry Professional | Accenture", 2)
    bullet(doc, "Develop and support SAP BTP administration workflows, including destinations, roles, identity, CTMS, CI/CD, and SAP Build Work Zone.")
    bullet(doc, "Apply SAP platform knowledge across SAP HANA, SAP Generative AI, and cloud administration tasks.")
    bullet(doc, "Bring an engineering and process-oriented approach to technical configuration, access flows, and delivery support.")

    heading(doc, "Process Engineer | SFA Semicon Philippines Corporation | Jun 2022 - Dec 2022", 2)
    bullet(doc, "Worked in a process engineering environment, applying manufacturing discipline, analysis, and operational problem-solving.")
    bullet(doc, "Built a practical foundation in systems thinking, documentation, and workflow improvement.")

    heading(doc, "Selected Projects", 1)
    heading(doc, "Obsidian - Skincare Ecommerce Website", 2)
    bullet(doc, "Built a luxury skincare ecommerce website currently under development with refined product sections, brand storytelling, reviews, cart/search navigation, and a trust-focused about page.")
    bullet(doc, "Tools and focus: Next.js, Supabase, ecommerce flows, UI design, brand-aware frontend development.")

    heading(doc, "Portfolio Website", 2)
    bullet(doc, "Created a black-and-white personal portfolio inspired by a compact personal index style, presenting SAP BTP work, certifications, web projects, and AI creative experiments.")

    heading(doc, "AI Video Creation", 2)
    bullet(doc, "Creates short-form AI-assisted videos, including food, skincare, and visual concept experiments.")

    heading(doc, "Certifications", 1)
    section_table(
        doc,
        [
            ("SAP", "SAP Certified - SAP BTP Administrator"),
            ("SAP", "SAP Certified - SAP Generative AI Developer"),
            ("Google / Coursera", "Foundations: Data, Data, Everywhere"),
        ],
    )

    heading(doc, "Education", 1)
    heading(doc, "Mapua University | B.S. Manufacturing Engineering | Jun 2015 - Aug 2021", 2)
    bullet(doc, "Engineering foundation covering manufacturing systems, process thinking, technical problem-solving, and professional discipline.")

    heading(doc, "Languages", 1)
    para(doc, "English | Filipino", size=10.5, color=BLACK, after=4)

    doc.save(OUT)


if __name__ == "__main__":
    create_doc()
    print(OUT)
